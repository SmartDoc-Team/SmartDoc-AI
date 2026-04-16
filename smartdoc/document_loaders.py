from __future__ import annotations

import base64
import io
import logging
import os
import re
import time
import unicodedata
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from zipfile import BadZipFile

import pdfplumber
from docx import Document as DocxDocument
from langchain_community.document_loaders import Docx2txtLoader
from pypdf import PdfReader
from pypdf.errors import PdfReadError

try:
    import fitz  # type: ignore[import-not-found]
except Exception:  # noqa: BLE001
    fitz = None

if fitz is None:
    try:
        import pymupdf as fitz  # type: ignore[import-not-found]
    except Exception:  # noqa: BLE001
        fitz = None

try:
    from PIL import Image
except Exception:  # noqa: BLE001
    Image = None

try:
    from pdf2image import convert_from_path
except Exception:  # noqa: BLE001
    convert_from_path = None

try:
    import pypdfium2 as pdfium
except Exception:  # noqa: BLE001
    pdfium = None

try:
    import pytesseract
except Exception:  # noqa: BLE001
    pytesseract = None

from smartdoc.config import settings
from smartdoc.ollama_client import OllamaClient


SUPPORTED_TYPES = {".pdf": "pdf", ".docx": "docx"}
CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")
SPACED_TOKEN_RE = re.compile(r"(?:\b[\wÀ-ỹ]\s+){3,}[\wÀ-ỹ]\b")
logger = logging.getLogger(__name__)


def _env_bool(name: str, default: bool = False) -> bool:
    raw_value = os.getenv(name)
    if raw_value is None:
        return default
    return raw_value.strip().lower() in {"1", "true", "yes", "y", "on"}


OCR_RENDER_DPI = int(os.getenv("OCR_RENDER_DPI", "150"))
OCR_MAX_IMAGE_SIDE = int(os.getenv("OCR_MAX_IMAGE_SIDE", "1600"))
OCR_JPEG_QUALITY = int(os.getenv("OCR_JPEG_QUALITY", "75"))
OCR_ENABLE_TESSERACT = _env_bool("OCR_ENABLE_TESSERACT", True)
OCR_TESSERACT_LANG = os.getenv("OCR_TESSERACT_LANG", "vie+eng")
OCR_TESSERACT_MIN_CHARS = int(os.getenv("OCR_TESSERACT_MIN_CHARS", "40"))
OCR_VISION_FALLBACK_MAX_PAGES = int(os.getenv("OCR_VISION_FALLBACK_MAX_PAGES", "3"))
OCR_DOCX_IMAGE_ENABLED = _env_bool("OCR_DOCX_IMAGE_ENABLED", True)
OCR_DOCX_IMAGE_MAX_ITEMS = int(os.getenv("OCR_DOCX_IMAGE_MAX_ITEMS", "8"))


ProgressCallback = Callable[[str], None]


def _report(progress_callback: ProgressCallback | None, message: str) -> None:
    logger.info(message)
    if progress_callback is None:
        return
    try:
        progress_callback(message)
    except Exception:  # noqa: BLE001
        logger.debug("Progress callback raised an exception", exc_info=True)


@dataclass(slots=True)  
class LoadedDocument:
    text: str | None
    pages: list[tuple[int, str]] | None
    source_name: str
    source_type: str


def normalize_text(text: str | bytes) -> str:
    if isinstance(text, bytes):
        text = text.decode("utf-8", errors="replace")

    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\ufeff", "").replace("\u200b", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
    text = CONTROL_CHARS_RE.sub(" ", text)

    def collapse_spaced_token(match: re.Match[str]) -> str:
        token = match.group(0)
        compact = re.sub(r"\s+", "", token)
        return compact if len(compact) >= 4 else token

    text = SPACED_TOKEN_RE.sub(collapse_spaced_token, text)
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _resolve_source_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[suffix]
    raise ValueError("Unsupported file type. Supported file types are: .pdf, .docx")


def load_document(file_path: str | Path, progress_callback: ProgressCallback | None = None) -> LoadedDocument:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.is_file():
        raise ValueError(f"Expected a file path, got: {path}")

    source_type = _resolve_source_type(path)
    _report(progress_callback, f"Start parsing {path.name} as {source_type.upper()}")

    try:
        if source_type == "pdf":
            pages = extract_pdf_pages(path, progress_callback=progress_callback)
            text = None
        else:
            _report(progress_callback, f"Extracting DOCX text from {path.name}")
            text = extract_docx_text(path, progress_callback=progress_callback)
            pages = None
    except Exception as exc:  
        raise ValueError(f"Failed to parse '{path.name}' as {source_type.upper()}: {exc}") from exc

    text = normalize_text(text) if text else None

    if not text and not pages:
        raise ValueError("The uploaded file did not produce readable text.")

    page_count = len(pages) if pages else 0
    _report(progress_callback, f"Completed parsing {path.name}: pages={page_count}, has_text={bool(text)}")

    return LoadedDocument(
        text=text,
        pages=pages,
        source_name=path.name,
        source_type=source_type,
    )


def _is_bad_pdf(pages: list[tuple[int, str]]) -> bool:
    if not pages:
        return True
    total_text = "".join(text for _, text in pages)
    # Nếu file là ảnh toàn bộ, số từ lấy ra được cực thấp (trung bình dưới 30 char mổi trang)
    if len(total_text) < len(pages) * 30:
        return True
        
    # Detect file lỗi font (chứa quá nhiều kí tự không map được cid: hoặc kí hiệu Unicode lỗi)
    garbage_patterns = [r"\(cid:\d+\)", r"\ufffd"]
    garbage_count = sum(len(re.findall(p, total_text)) for p in garbage_patterns)
    if len(total_text) > 0 and (garbage_count / len(total_text)) > 0.05:
        return True
        
    return False


def _render_pdf_images_with_pymupdf(path: Path, dpi: int = 200) -> list[object]:
    if fitz is None:
        raise ValueError("PyMuPDF is not installed.")
    if Image is None:
        raise ValueError("Pillow is not installed.")

    images: list[object] = []
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)

    with fitz.open(str(path)) as pdf_doc:
        for page in pdf_doc:
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            images.append(image)

    if not images:
        raise ValueError("No pages were rendered from PDF using PyMuPDF.")
    return images


def _render_pdf_images_with_pdf2image(path: Path) -> list[object]:
    if convert_from_path is None:
        raise ValueError("pdf2image is not installed.")
    try:
        return convert_from_path(str(path))
    except Exception as exc:
        raise ValueError(f"Failed to convert PDF to images with pdf2image: {exc}") from exc


def _render_pdf_images_with_pypdfium2(path: Path, dpi: int = 200) -> list[object]:
    if pdfium is None:
        raise ValueError("pypdfium2 is not installed.")
    if Image is None:
        raise ValueError("Pillow is not installed.")

    images: list[object] = []
    scale = dpi / 72.0

    pdf_doc = pdfium.PdfDocument(str(path))
    try:
        for page_index in range(len(pdf_doc)):
            page = pdf_doc[page_index]
            try:
                bitmap = page.render(scale=scale)
                try:
                    image = bitmap.to_pil().copy()
                    images.append(image)
                finally:
                    bitmap.close()
            finally:
                page.close()
    finally:
        pdf_doc.close()

    if not images:
        raise ValueError("No pages were rendered from PDF using pypdfium2.")
    return images


def _prepare_ocr_image(img: object, max_side: int) -> object:
    if Image is None:
        return img

    width, height = img.size
    largest_side = max(width, height)
    if largest_side <= max_side:
        return img

    scale = max_side / float(largest_side)
    new_size = (max(1, int(width * scale)), max(1, int(height * scale)))
    return img.resize(new_size, Image.Resampling.LANCZOS)


def _render_pdf_images(path: Path, progress_callback: ProgressCallback | None = None) -> list[object]:
    errors: list[str] = []
    images: list[object] = []

    try:
        images = _render_pdf_images_with_pymupdf(path, dpi=OCR_RENDER_DPI)
        _report(progress_callback, "Rendered PDF pages with PyMuPDF")
    except Exception as exc:  # noqa: BLE001
        errors.append(f"pymupdf: {exc}")

    if not images:
        try:
            images = _render_pdf_images_with_pypdfium2(path, dpi=OCR_RENDER_DPI)
            _report(progress_callback, "Rendered PDF pages with pypdfium2")
        except Exception as exc:  # noqa: BLE001
            errors.append(f"pypdfium2: {exc}")

    if not images:
        try:
            images = _render_pdf_images_with_pdf2image(path)
            _report(progress_callback, "Rendered PDF pages with pdf2image")
        except Exception as exc:  # noqa: BLE001
            errors.append(f"pdf2image: {exc}")

    if not images:
        raise ValueError("Failed to convert PDF to images for OCR: " + "; ".join(errors))

    return images


def _is_low_quality_ocr_result(text: str, min_chars: int = OCR_TESSERACT_MIN_CHARS) -> bool:
    normalized = normalize_text(text)
    if len(normalized) < min_chars:
        return True

    alnum_count = sum(1 for char in normalized if char.isalnum())
    if alnum_count / max(1, len(normalized)) < 0.30:
        return True

    return False


def _extract_images_with_tesseract_plus_selective_vision(
    images: list[object],
    owner_path: Path,
    progress_callback: ProgressCallback | None = None,
    image_label: str = "image",
) -> list[tuple[int, str]]:
    if not OCR_ENABLE_TESSERACT:
        raise ValueError("Tesseract OCR is disabled by OCR_ENABLE_TESSERACT.")
    if pytesseract is None:
        raise ValueError("pytesseract is not installed.")

    total_images = len(images)
    _report(progress_callback, f"Starting fast OCR with Tesseract on {total_images} {image_label}(s)")

    pages: list[tuple[int, str]] = []
    low_quality_pages: list[int] = []

    for page_number, img in enumerate(images, start=1):
        prepared_img = _prepare_ocr_image(img, max_side=OCR_MAX_IMAGE_SIDE)
        extracted = pytesseract.image_to_string(prepared_img, lang=OCR_TESSERACT_LANG, config="--oem 1 --psm 6")
        normalized = normalize_text(extracted)
        pages.append((page_number, normalized))

        if _is_low_quality_ocr_result(normalized):
            low_quality_pages.append(page_number)

    if not low_quality_pages:
        _report(progress_callback, f"Tesseract OCR succeeded for all {image_label}(s)")
        return [(page_num, text) for page_num, text in pages if text]

    _report(progress_callback, f"Tesseract low-quality {image_label}(s): {len(low_quality_pages)}")

    selected_for_vision = low_quality_pages[: max(0, OCR_VISION_FALLBACK_MAX_PAGES)]
    if selected_for_vision:
        _report(
            progress_callback,
            f"Running selective vision OCR for {len(selected_for_vision)} {image_label}(s): {selected_for_vision}",
        )
        vision_pages = _extract_pdf_pages_with_vision(
            owner_path,
            progress_callback=progress_callback,
            pre_rendered_images=images,
            page_numbers=selected_for_vision,
        )
        vision_map = {page_num: text for page_num, text in vision_pages}
        pages = [(page_num, vision_map.get(page_num, text)) for page_num, text in pages]

    return [(page_num, text) for page_num, text in pages if text]


def _extract_pdf_pages_with_vision(
    path: Path,
    progress_callback: ProgressCallback | None = None,
    pre_rendered_images: list[object] | None = None,
    page_numbers: list[int] | None = None,
) -> list[tuple[int, str]]:
    images = pre_rendered_images or _render_pdf_images(path, progress_callback=progress_callback)

    vision_model = os.getenv("OLLAMA_VISION_MODEL", "llava")
    
    # Thiết lập LlaVA client
    client = OllamaClient(
        base_url=settings.ollama_base_url,
        model=vision_model,
        temperature=settings.ollama_temperature
    )
    pages: list[tuple[int, str]] = []
    selected_pages = page_numbers or list(range(1, len(images) + 1))
    total_pages = len(selected_pages)
    _report(progress_callback, f"Starting vision OCR with model '{vision_model}' on {total_pages} page(s)")
    
    prompt = (
        "You are an expert OCR AI. Extract ALL the text from this document image accurately. "
        "Preserve the original language (e.g. Vietnamese). Convert any tables into Markdown table format. "
        "Do NOT add any conversational filler, intro or outro. Output only the extracted text."
    )

    for position, page_number in enumerate(selected_pages, start=1):
        if page_number < 1 or page_number > len(images):
            continue

        img = images[page_number - 1]
        started_at = time.perf_counter()
        _report(progress_callback, f"OCR page {position}/{total_pages} (doc page {page_number})")

        prepared_img = _prepare_ocr_image(img, max_side=OCR_MAX_IMAGE_SIDE)
        if prepared_img.size != img.size:
            _report(
                progress_callback,
                f"Resized page {page_number}: {img.size[0]}x{img.size[1]} -> {prepared_img.size[0]}x{prepared_img.size[1]}",
            )

        buffered = io.BytesIO()
        prepared_img.save(buffered, format="JPEG", quality=OCR_JPEG_QUALITY, optimize=True)
        img_b64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
        
        try:
            extracted_text = client.generate(prompt=prompt, images=[img_b64])
            pages.append((page_number, extracted_text))
            elapsed = time.perf_counter() - started_at
            _report(progress_callback, f"OCR done page {position}/{total_pages} (doc page {page_number}) in {elapsed:.1f}s")
        except Exception as e:
            pages.append((page_number, f"[Vision OCR Error on page {page_number}: {e}]"))
            _report(progress_callback, f"OCR failed on doc page {page_number}: {e}")
            
    _report(progress_callback, "Finished vision OCR phase")
    return pages


def _extract_pdf_pages_with_tesseract_plus_selective_vision(
    path: Path,
    progress_callback: ProgressCallback | None = None,
) -> list[tuple[int, str]]:
    images = _render_pdf_images(path, progress_callback=progress_callback)
    return _extract_images_with_tesseract_plus_selective_vision(
        images=images,
        owner_path=path,
        progress_callback=progress_callback,
        image_label="page",
    )

def extract_pdf_pages(path: Path, progress_callback: ProgressCallback | None = None) -> list[tuple[int, str]]:
    errors: list[str] = []
    pages: list[tuple[int, str]] = []
    _report(progress_callback, "Trying pypdf text extraction")

    try:
        pages = _extract_pdf_pages_with_pypdf(path)
        if not _is_bad_pdf(pages):
            _report(progress_callback, "pypdf extraction succeeded")
            return pages
        _report(progress_callback, "pypdf output looks low quality, trying fallback")
    except (PdfReadError, OSError, ValueError) as exc:
        errors.append(f"pypdf: {exc}")
        _report(progress_callback, f"pypdf failed: {exc}")

    _report(progress_callback, "Trying pdfplumber text extraction")
    try:
        plumber_pages = _extract_pdf_pages_with_pdfplumber(path)
        if not _is_bad_pdf(plumber_pages):
            _report(progress_callback, "pdfplumber extraction succeeded")
            return plumber_pages
        if not pages:
            pages = plumber_pages
        _report(progress_callback, "pdfplumber output looks low quality, trying OCR")
    except Exception as exc:  
        errors.append(f"pdfplumber: {exc}")
        _report(progress_callback, f"pdfplumber failed: {exc}")

    _report(progress_callback, "Trying fast OCR (Tesseract) before full vision OCR")
    try:
        fast_ocr_pages = _extract_pdf_pages_with_tesseract_plus_selective_vision(path, progress_callback=progress_callback)
        if fast_ocr_pages and not _is_bad_pdf(fast_ocr_pages):
            _report(progress_callback, "Fast OCR extraction succeeded")
            return fast_ocr_pages
        if fast_ocr_pages:
            pages = fast_ocr_pages
        _report(progress_callback, "Fast OCR output still looks low quality, trying full vision OCR")
    except Exception as exc:
        errors.append(f"tesseract_ocr: {exc}")
        _report(progress_callback, f"Fast OCR failed: {exc}")

    # Fallback full LLaVA OCR nếu text trích xuất báo lỗi Bad PDF.
    try:
        vision_pages = _extract_pdf_pages_with_vision(path, progress_callback=progress_callback)
        if vision_pages:
            _report(progress_callback, "Vision OCR extraction succeeded")
            return vision_pages
    except Exception as exc:
        errors.append(f"vision_ocr: {exc}")
        _report(progress_callback, f"Vision OCR failed: {exc}")

    # Nếu OCR Vision cũng lỗi, trả về cái ít tệ nhất đã parse được
    if pages:
        _report(progress_callback, "Returning best-effort pages from previous extraction")
        return pages

    if errors:
        raise RuntimeError("; ".join(errors))
    raise ValueError("No readable text was extracted from the PDF file.")


def _extract_pdf_pages_with_pypdf(path: Path) -> list[tuple[int, str]]:
    reader = PdfReader(str(path))
    pages: list[tuple[int, str]] = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = normalize_text(page.extract_text() or "")
        if page_text:
            pages.append((page_number, page_text))

    if not pages:
        raise ValueError("pypdf could not extract readable text.")

    return pages


def _extract_pdf_pages_with_pdfplumber(path: Path) -> list[tuple[int, str]]:
    pages: list[tuple[int, str]] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = normalize_text(page.extract_text() or "")
            if page_text:
                pages.append((page_number, page_text))

    if not pages:
        raise ValueError("pdfplumber could not extract readable text.")

    return pages


def extract_docx_text(path: Path, progress_callback: ProgressCallback | None = None) -> str:
    errors: list[str] = []
    candidates: list[str] = []

    _report(progress_callback, "Trying python-docx text extraction")

    try:
        text = _extract_docx_text_with_python_docx(path)
        if text:
            candidates.append(text)
    except (BadZipFile, OSError, ValueError) as exc:
        errors.append(f"python-docx: {exc}")
    except Exception as exc:  
        errors.append(f"python-docx: {exc}")

    _report(progress_callback, "Trying Docx2txtLoader extraction")
    try:
        text = _extract_docx_text_with_loader(path)
        if text:
            candidates.append(text)
    except Exception as exc: 
        errors.append(f"Docx2txtLoader: {exc}")

    best_text = max(candidates, key=len, default="")
    if best_text and not _is_low_quality_ocr_result(best_text):
        _report(progress_callback, "DOCX text extraction succeeded")
        return best_text

    image_text = ""
    if OCR_DOCX_IMAGE_ENABLED:
        _report(progress_callback, "DOCX text looks weak, trying OCR on embedded images")
        try:
            image_text = _extract_docx_embedded_image_text(path, progress_callback=progress_callback)
        except Exception as exc:  # noqa: BLE001
            errors.append(f"docx_image_ocr: {exc}")

    if best_text and image_text:
        return normalize_text(f"{best_text}\n\n{image_text}")
    if image_text:
        return image_text
    if best_text:
        return best_text

    if errors:
        raise RuntimeError("; ".join(errors))
    raise ValueError("No readable text was extracted from the DOCX file.")


def _extract_docx_embedded_image_text(path: Path, progress_callback: ProgressCallback | None = None) -> str:
    if Image is None:
        raise ValueError("Pillow is not installed.")

    doc = DocxDocument(str(path))
    images: list[object] = []

    for relationship in doc.part.rels.values():
        rel_type = getattr(relationship, "reltype", "")
        if "image" not in rel_type:
            continue
        target = getattr(relationship, "target_part", None)
        blob = getattr(target, "blob", None)
        if not blob:
            continue
        with Image.open(io.BytesIO(blob)) as pil_image:
            images.append(pil_image.convert("RGB"))

    if not images:
        raise ValueError("No embedded images found in DOCX for OCR fallback.")

    if OCR_DOCX_IMAGE_MAX_ITEMS > 0:
        images = images[:OCR_DOCX_IMAGE_MAX_ITEMS]

    pages = _extract_images_with_tesseract_plus_selective_vision(
        images=images,
        owner_path=path,
        progress_callback=progress_callback,
        image_label="docx image",
    )
    texts = [text for _, text in pages if text]
    if not texts:
        raise ValueError("OCR could not extract text from DOCX embedded images.")

    return normalize_text("\n\n".join(texts))


def _extract_docx_text_with_python_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        heading_level = _heading_level_from_style(style_name)
        if heading_level:
            parts.append(f"{'#' * heading_level} {text}")
        else:
            parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        markdown_rows = _table_to_markdown_rows(table)
        if markdown_rows:
            parts.append(f"[Table {table_index}]")
            parts.extend(markdown_rows)

    return "\n\n".join(parts)


def _extract_docx_text_with_loader(path: Path) -> str:
    loader = Docx2txtLoader(str(path))
    documents = loader.load()
    parts: list[str] = []

    for doc in documents:
        page_content = normalize_text(doc.page_content)
        if page_content:
            parts.append(page_content)

    if not parts:
        raise ValueError("Docx2txtLoader returned empty content.")

    return "\n\n".join(parts)


def _heading_level_from_style(style_name: str) -> int | None:
    match = re.search(r"heading\s*(\d+)", style_name.lower())
    if not match:
        return None
    level = int(match.group(1))
    return max(1, min(level, 6))


def _table_to_markdown_rows(table: object) -> list[str]:
    rows: list[list[str]] = []

    for row in table.rows:  # type: ignore[attr-defined]
        cells = [normalize_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append([cell if cell else "-" for cell in cells])

    if not rows:
        return []

    max_columns = max(len(row) for row in rows)

    normalized_rows: list[list[str]] = []
    for row in rows:
        if len(row) < max_columns:
            row = row + ["-"] * (max_columns - len(row))
        normalized_rows.append(row)

    markdown_rows: list[str] = [" | ".join(normalized_rows[0])]
    if len(normalized_rows) > 1:
        markdown_rows.append(" | ".join(["---"] * max_columns))
        markdown_rows.extend(" | ".join(row) for row in normalized_rows[1:])

    return markdown_rows
